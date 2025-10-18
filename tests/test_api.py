import io
import json
import os
import re
import subprocess
import time
from typing import Dict, Tuple

import pytest
import requests
from dotenv import load_dotenv

load_dotenv()

# ---------- Константы окружения ----------
URL_PREFIX = os.getenv("AUTO_SUMMARIZATION_URL_PREFIX", "/v1").rstrip("/")
API_PORT = int(os.getenv("AUTO_SUMMARIZATION_API_PORT", "8000"))
API_HOST = os.getenv("AUTO_SUMMARIZATION_API_HOST", "0.0.0.0")
BASE_URL = f"http://{API_HOST}:{API_PORT}"
MAX_TEXT_LEN = int(os.getenv("AUTO_SUMMARIZATION_MAX_TEXT_LENGTH", "100000"))
SUPPORTED_FORMATS = tuple(
    s.strip().lower()
    for s in os.getenv("AUTO_SUMMARIZATION_SUPPORTED_FORMATS", "txt,doc,docx,pdf,odt").split(",")
    if s.strip()
)

# DEBUG=1 → сервис ожидает заголовок user_id, иначе Authorization
AUTH_HEADER_NAME = "user_id" if int(os.getenv("DEBUG", "0")) != 0 else "Authorization"

COMPOSE_FILE = "docker-compose.yml"
DEV_DIR = "dev"
LOG_PATH = os.path.join(DEV_DIR, "content.txt")


def _auth_headers(user_id: str | None) -> Dict[str, str]:
    if user_id is None:
        return {}
    return {AUTH_HEADER_NAME: user_id}


def _wait_healthy(timeout: int = 90):
    deadline = time.time() + timeout
    last_error = None
    while time.time() < deadline:
        try:
            r = requests.get(f"{BASE_URL}/health", timeout=3)
            if r.status_code == 200 and r.json().get("status") == "ok":
                return
        except Exception as e:
            last_error = e
        time.sleep(1)
    raise RuntimeError(f"Service is not healthy: {last_error}")


@pytest.mark.asyncio
class TestAPI:
    @classmethod
    def setup_class(cls):
        os.makedirs(DEV_DIR, exist_ok=True)
        # поднимаем стек
        with open(LOG_PATH, "w", encoding="utf-8") as f:
            f.write("**Logs**\n\n")
        with open(LOG_PATH, "a", encoding="utf-8") as f:
            subprocess.run(
                ["docker", "compose", "-f", COMPOSE_FILE, "up", "--build", "-d"],
                stdout=f, stderr=f, check=False
            )
        _wait_healthy(timeout=120)

    @classmethod
    def teardown_class(cls):
        # останавливаем стек (без ошибок, чтобы не падать из-за already down)
        with open(LOG_PATH, "a", encoding="utf-8") as f:
            subprocess.run(
                ["docker", "compose", "-f", COMPOSE_FILE, "down", "-v"],
                stdout=f, stderr=f, check=False
            )

    # --------- /health ----------
    def test_health(self):
        r = requests.get(f"{BASE_URL}/health", timeout=5)
        assert r.status_code == 200
        assert r.json() == {"status": "ok"}

    # --------- /v1/user/* ----------
    def test_users__create_list_delete(self):
        user_id = "u_test_users_flow"
        # create
        r = requests.post(
            f"{BASE_URL}{URL_PREFIX}/user/create_user",
            json={"user_id": user_id, "temporary": False},
            timeout=10,
        )
        assert r.status_code == 200
        assert r.json()["status"] in ("created", "exist")

        # idempotent create → exist
        r2 = requests.post(
            f"{BASE_URL}{URL_PREFIX}/user/create_user",
            json={"user_id": user_id, "temporary": False},
            timeout=10,
        )
        assert r2.status_code == 200
        assert r2.json()["status"] == "exist"

        # list (содержит только не temporary; мы создаём non-temp)
        r3 = requests.get(f"{BASE_URL}{URL_PREFIX}/user/get_users", timeout=10)
        assert r3.status_code == 200
        users = r3.json()["users"]
        assert any(u["user_id"] == user_id for u in users)

        # delete
        r4 = requests.delete(
            f"{BASE_URL}{URL_PREFIX}/user/delete_user",
            json={"user_id": user_id},
            timeout=10,
        )
        assert r4.status_code == 200
        assert r4.json()["status"] in ("deleted",)

        # delete non-existing → not_found
        r5 = requests.delete(
            f"{BASE_URL}{URL_PREFIX}/user/delete_user",
            json={"user_id": user_id},
            timeout=10,
        )
        assert r5.status_code == 200
        assert r5.json()["status"] in ("not_found",)

    # --------- /v1/analysis/* ----------
    def test_analysis__analyze_types_positive(self):
        r = requests.get(f"{BASE_URL}{URL_PREFIX}/analysis/analyze_types", timeout=10)
        assert r.status_code == 200
        data = r.json()
        assert isinstance(data["categories"], list) and len(data["categories"]) > 0
        assert isinstance(data["choices"], list) and len(data["choices"]) > 0

    def test_analysis__load_document_txt_positive(self):
        # простой txt
        files = {"document": ("note.txt", b"Hello\nWorld", "text/plain")}
        r = requests.post(f"{BASE_URL}{URL_PREFIX}/analysis/load_document", files=files, timeout=10)
        assert r.status_code == 200
        assert "Hello" in r.json()["text"]

    def test_analysis__load_document_docx_positive(self):
        # генерируем минимальный docx в памяти
        try:
            from docx import Document  # type: ignore
        except Exception:
            pytest.skip("python-docx не установлен в окружении тестов")

        buf = io.BytesIO()
        d = Document()
        d.add_paragraph("Docx Line 1")
        d.add_paragraph("Docx Line 2")
        d.save(buf)
        buf.seek(0)

        files = {"document": ("file.docx", buf.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = requests.post(f"{BASE_URL}{URL_PREFIX}/analysis/load_document", files=files, timeout=10)
        assert r.status_code == 200
        text = r.json()["text"]
        assert "Docx Line 1" in text and "Docx Line 2" in text

    def test_analysis__load_document_unsupported_negative(self):
        files = {"document": ("binary.xyz", b"\x00\x01\x02", "application/octet-stream")}
        r = requests.post(f"{BASE_URL}{URL_PREFIX}/analysis/load_document", files=files, timeout=10)
        assert r.status_code == 400
        assert r.json()["detail"] == "Unsupported document format"

    # --------- /v1/chat_session/* ----------
    def _create_user_session(
        self, user_id: str, text: str, title: str = "", category_index: int = 0, choices: Tuple[int, ...] = ()
    ) -> Tuple[str, Dict]:
        # создаём сессию с choices, которых нет в шаблонах (чтобы не дергать LLM/Transformers)
        payload = {
            "title": title,
            "text": text,
            "category": category_index,
            "choices": list(choices),
            "temporary": False,
        }
        r = requests.post(
            f"{BASE_URL}{URL_PREFIX}/chat_session/create",
            json=payload,
            headers=_auth_headers(user_id),
            timeout=20,
        )
        assert r.status_code == 200, r.text
        data = r.json()
        assert data["session_id"]
        return data["session_id"], data

    def test_session__fetch_page_requires_auth_negative(self):
        r = requests.get(f"{BASE_URL}{URL_PREFIX}/chat_session/fetch_page", timeout=10)
        assert r.status_code == 400
        assert "Authorization header is required" in r.text or "Bad Request" in r.text

    def test_session__create_and_fetch_page_positive(self):
        user_id = "u_create_fetch"
        # fetch_page до создания → пусто
        r0 = requests.get(
            f"{BASE_URL}{URL_PREFIX}/chat_session/fetch_page",
            headers=_auth_headers(user_id),
            timeout=10,
        )
        assert r0.status_code == 200
        assert r0.json()["sessions"] == []

        # создаём
        session_id, _ = self._create_user_session(
            user_id=user_id,
            text="Классический пример текста об экономике и рынках",
            title="",
            category_index=0,
            choices=(9999,),  # отсутствующий индекс → без LLM
        )

        # fetch_page → содержит созданную
        r1 = requests.get(
            f"{BASE_URL}{URL_PREFIX}/chat_session/fetch_page",
            headers=_auth_headers(user_id),
            timeout=10,
        )
        assert r1.status_code == 200
        sessions = r1.json()["sessions"]
        assert any(s["session_id"] == session_id for s in sessions)

    def test_session__create_invalid_category_negative(self):
        user_id = "u_invalid_category"
        r = requests.post(
            f"{BASE_URL}{URL_PREFIX}/chat_session/create",
            json={
                "title": "",
                "text": "text",
                "category": 9999,       # несуществующая категория
                "choices": [0, 1, 2],
                "temporary": False,
            },
            headers=_auth_headers(user_id),
            timeout=10,
        )
        assert r.status_code == 400
        assert r.json()["detail"] in ("Invalid category index",)

    def test_session__create_text_length_exceeded_negative(self):
        user_id = "u_text_len"
        huge_text = "a" * (MAX_TEXT_LEN + 1)
        r = requests.post(
            f"{BASE_URL}{URL_PREFIX}/chat_session/create",
            json={"title": "", "text": huge_text, "category": 0, "choices": [], "temporary": False},
            headers=_auth_headers(user_id),
            timeout=20,
        )
        assert r.status_code == 400
        assert f"Длина одного документа превышает лимит {MAX_TEXT_LEN} символов" in r.text

    def test_session__info_update_title_update_summarization_flow(self):
        user_id = "u_update_flow"
        text = "Небольшой текст для сессии"
        session_id, created = self._create_user_session(
            user_id=user_id, text=text, category_index=0, choices=()
        )

        # получаем полную информацию
        r_info = requests.get(
            f"{BASE_URL}{URL_PREFIX}/chat_session/{session_id}",
            headers=_auth_headers(user_id),
            timeout=10,
        )
        assert r_info.status_code == 200
        info = r_info.json()
        assert info["session_id"] == session_id
        version0 = info["version"]

        # update_title (OK)
        r_title = requests.post(
            f"{BASE_URL}{URL_PREFIX}/chat_session/update_title",
            json={"session_id": session_id, "title": "Новый заголовок", "version": version0},
            headers=_auth_headers(user_id),
            timeout=10,
        )
        assert r_title.status_code == 200
        info2 = r_title.json()
        assert info2["title"] == "Новый заголовок"
        version1 = info2["version"]

        # update_title с неверной версией → 400
        r_title_bad = requests.post(
            f"{BASE_URL}{URL_PREFIX}/chat_session/update_title",
            json={"session_id": session_id, "title": "Ещё заголовок", "version": version0},
            headers=_auth_headers(user_id),
            timeout=10,
        )
        assert r_title_bad.status_code == 400
        assert r_title_bad.json()["detail"] == "Version mismatch"

        # update_summarization (OK, choices отсутствуют → без LLM)
        r_sum = requests.post(
            f"{BASE_URL}{URL_PREFIX}/chat_session/update_summarization",
            json={
                "session_id": session_id,
                "text": text + " + дополнение",
                "category": 0,
                "choices": [424242],  # отсутствующий индекс
                "version": version1,
            },
            headers=_auth_headers(user_id),
            timeout=20,
        )
        assert r_sum.status_code == 200
        data_sum = r_sum.json()
        assert "content" in data_sum and isinstance(data_sum["content"], dict)

        # update_summarization с неверной версией → 400
        r_sum_bad = requests.post(
            f"{BASE_URL}{URL_PREFIX}/chat_session/update_summarization",
            json={
                "session_id": session_id,
                "text": text,
                "category": 0,
                "choices": [],
                "version": version1,  # уже устарел
            },
            headers=_auth_headers(user_id),
            timeout=10,
        )
        assert r_sum_bad.status_code == 400
        assert r_sum_bad.json()["detail"] == "Version mismatch"

    def test_session__search_positive_and_negatives(self):
        user_id = "u_search"
        # создаём пару сессий
        sid1, _ = self._create_user_session(
            user_id=user_id, text="Рынки растут. Акции X увеличились.", category_index=0, choices=(999,)
        )
        sid2, _ = self._create_user_session(
            user_id=user_id, text="Футбол: команда Y победила со счётом 2-1.", category_index=0, choices=(999,)
        )

        # Positive: поиск по слову "Акции" → должен найти sid1
        r_ok = requests.get(
            f"{BASE_URL}{URL_PREFIX}/chat_session/search",
            params={"query": "Акции"},
            headers=_auth_headers(user_id),
            timeout=10,
        )
        assert r_ok.status_code == 200
        results = r_ok.json()["sessions"]
        assert any(item["session_id"] == sid1 for item in results)

        # Negative: пустой query → 422 (валидация FastAPI на min_length=1)
        r_422 = requests.get(
            f"{BASE_URL}{URL_PREFIX}/chat_session/search",
            params={"query": ""},
            headers=_auth_headers(user_id),
            timeout=10,
        )
        assert r_422.status_code == 422

        # Negative: без заголовка авторизации
        r_400 = requests.get(
            f"{BASE_URL}{URL_PREFIX}/chat_session/search",
            params={"query": "anything"},
            timeout=10,
        )
        assert r_400.status_code == 400

    def test_session__download_pdf_positive_and_not_found(self):
        user_id = "u_download"
        sid, _ = self._create_user_session(
            user_id=user_id, text="Текст для экспорта в PDF", category_index=0, choices=()
        )

        # JSON-ответ (base64) через Accept: application/json
        r_json = requests.get(
            f"{BASE_URL}{URL_PREFIX}/chat_session/download/{sid}/pdf",
            headers={**_auth_headers(user_id), "Accept": "application/json"},
            timeout=20,
        )
        assert r_json.status_code == 200
        payload = r_json.json()
        assert set(payload.keys()) == {"filename", "content_type", "data"}
        assert payload["filename"].endswith(".pdf")
        assert payload["content_type"] in ("application/pdf", "application/octet-stream")
        assert isinstance(payload["data"], str) and len(payload["data"]) > 0

        # not found для несуществующей сессии
        r_nf = requests.get(
            f"{BASE_URL}{URL_PREFIX}/chat_session/download/does-not-exist/pdf",
            headers={**_auth_headers(user_id), "Accept": "application/json"},
            timeout=10,
        )
        assert r_nf.status_code in (404,)

    def test_session__delete_positive_and_second_call_error_status(self):
        user_id = "u_delete"
        sid, _ = self._create_user_session(
            user_id=user_id, text="Удаляемая сессия", category_index=0, choices=()
        )
        # Удаляем
        r_del = requests.delete(
            f"{BASE_URL}{URL_PREFIX}/chat_session/delete",
            json={"session_id": sid},
            headers=_auth_headers(user_id),
            timeout=10,
        )
        assert r_del.status_code == 200
        assert r_del.json()["status"] in ("SUCCESS",)

        # Повторное удаление → SUCCESS не будет, ожидаем ERROR
        r_del2 = requests.delete(
            f"{BASE_URL}{URL_PREFIX}/chat_session/delete",
            json={"session_id": sid},
            headers=_auth_headers(user_id),
            timeout=10,
        )
        assert r_del2.status_code == 200
        assert r_del2.json()["status"] in ("ERROR",)
